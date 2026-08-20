"""Multi-engine text extraction.

We run more than one engine on purpose. Engines from different families disagree about what a
PDF contains, and that disagreement is itself a detector - see divergence.py. Adapted from
ats-extraxt-test's engine registry, trimmed to the four that are actually installed here.
"""

from __future__ import annotations

from pathlib import Path

# pymupdf and pypdfium2 share a "render-oriented" view and drop off-page content; pdfplumber
# and pdfminer share a "content-stream" view and keep it. One from each family is the minimum
# that makes divergence meaningful.
FAMILY = {"pymupdf": "render", "pypdfium2": "render", "pdfplumber": "stream", "pdfminer": "stream"}
PRIMARY = "pymupdf"


def _pymupdf(path: str) -> str:
    import pymupdf

    with pymupdf.open(path) as doc:
        return "\n".join(p.get_text() for p in doc)


def _pypdfium2(path: str) -> str:
    import pypdfium2

    doc = pypdfium2.PdfDocument(path)
    try:
        return "\n".join(doc[i].get_textpage().get_text_bounded() for i in range(len(doc)))
    finally:
        doc.close()


def _pdfplumber(path: str) -> str:
    import pdfplumber

    with pdfplumber.open(path) as doc:
        return "\n".join(p.extract_text() or "" for p in doc.pages)


def _pdfminer(path: str) -> str:
    from pdfminer.high_level import extract_text

    return extract_text(path)


def _docx(path: str) -> str:
    import docx

    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    parts += [c.text for t in d.tables for r in t.rows for c in r.cells]
    return "\n".join(parts)


ENGINES = {
    "pymupdf": _pymupdf,
    "pypdfium2": _pypdfium2,
    "pdfplumber": _pdfplumber,
    "pdfminer": _pdfminer,
}


def extract_all(path: str | Path) -> dict[str, str]:
    """Every engine's view of the file. A failing engine yields "" rather than raising."""
    path = str(path)
    if path.lower().endswith(".docx"):
        return {"docx": _docx(path)}
    if path.lower().endswith(".txt"):
        return {"txt": Path(path).read_text(errors="replace")}
    out = {}
    for name, fn in ENGINES.items():
        try:
            out[name] = fn(path)
        except Exception:
            out[name] = ""
    return out


def primary_text(views: dict[str, str]) -> str:
    for name in (PRIMARY, "docx", "txt", "pdfplumber", "pdfminer", "pypdfium2"):
        if views.get(name):
            return views[name]
    return next((v for v in views.values() if v), "")
